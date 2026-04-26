"""Report renderer + ReportContext tests.

The renderer tests need ``docxtpl`` installed; they skip cleanly when
it isn't, the same way the pytest-qt-dependent suites do. The
context-builder tests don't need docxtpl and run unconditionally — a
context can be assembled and inspected without ever rendering a
template.
"""

from __future__ import annotations

import json
from datetime import UTC, datetime
from typing import TYPE_CHECKING

import pytest

from caseforge.model import (
    Case,
    CustodyRecord,
    ExaminerIdentity,
    ExamScope,
)
from caseforge.report.context import build_context
from caseforge.report.suggestions_reader import (
    suggestions_path as caseguide_suggestions_path,
)
from caseforge.storage import write_case

if TYPE_CHECKING:
    from pathlib import Path


_NOW = datetime(2026, 4, 26, 12, 0, tzinfo=UTC)


def _seed_case(case_dir: Path) -> Case:
    case = Case(
        name="Smoketest CSAM",
        case_reference="ST-001",
        created_at=_NOW,
        updated_at=_NOW,
        examiner=ExaminerIdentity(
            name="Alex Smith",
            organisation="Cyber Crimes Unit",
            badge_id="CCU-0421",
        ),
        scope=ExamScope(
            exam_type="CSAM possession",
            primary_tool="axiom",
            device_classes=["windows-laptop"],
            evidence_items=["E01 image"],
            agencies=["FBI", "ICAC"],
            summary="Single laptop, single image, single user.",
        ),
        custody=CustodyRecord(
            received_at=_NOW,
            received_from="Det. Rivers",
            delivery_method="in person",
            evidence_bag_ids=["BAG-12", "BAG-13"],
        ),
    )
    case_dir.mkdir(parents=True, exist_ok=True)
    write_case(case_dir, case)
    return case


def _write_suggestions(case_dir: Path) -> None:
    target = caseguide_suggestions_path(case_dir)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(
        json.dumps(
            {
                "schema_version": 2,
                "generated_at": "2026-04-26T11:00:00+00:00",
                "scope_summary": "CSAM possession; Win11 laptop.",
                "playbooks": ["verify-image-hash", "axiom-ci-processing"],
                "suggestions": [
                    {
                        "id": "verify-image-hash",
                        "action": "Verify SHA-256.",
                        "priority": "required",
                        "category": "verification",
                        "expected_result": "Hash matches.",
                        "completed": True,
                        "completed_at": "2026-04-26T11:30:00+00:00",
                    },
                    {
                        "id": "axiom-ci-processing",
                        "action": "Run AXIOM Process.",
                        "priority": "required",
                        "category": "processing",
                        "depends_on": ["verify-image-hash"],
                    },
                ],
            }
        ),
        encoding="utf-8",
    )


# ----------------------------------------------------- context builder


def test_context_builds_from_case_only(tmp_path: Path) -> None:
    """A fresh case dir with no suggestions / sessions still builds a context."""
    case_dir = tmp_path / "case"
    _seed_case(case_dir)

    ctx = build_context(case_dir, now=_NOW)
    assert ctx.case.name == "Smoketest CSAM"
    assert ctx.case.reference == "ST-001"
    assert ctx.examiner.name == "Alex Smith"
    assert ctx.examiner.is_present is True
    assert ctx.scope.exam_type == "CSAM possession"
    assert ctx.scope.evidence_items_csv == "E01 image"
    assert ctx.suggestions.has_data is False
    assert ctx.suggestions.total == 0
    assert ctx.suggestions.completed == []
    assert ctx.sessions.count == 0
    assert ctx.generated_at_iso == _NOW.isoformat()


def test_context_picks_up_suggestions_when_present(tmp_path: Path) -> None:
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    _write_suggestions(case_dir)

    ctx = build_context(case_dir, now=_NOW)
    assert ctx.suggestions.has_data is True
    assert ctx.suggestions.total == 2
    assert ctx.suggestions.completed_count == 1
    assert ctx.suggestions.required_count == 2
    assert [s.id for s in ctx.suggestions.completed] == ["verify-image-hash"]
    assert [s.id for s in ctx.suggestions.open] == ["axiom-ci-processing"]
    assert ctx.suggestions.completed[0].completed_at_str.startswith("2026-04-26 11:30")


def test_context_template_dict_is_render_safe(tmp_path: Path) -> None:
    """The flattened dict must not contain dataclass instances at any depth.

    docxtpl's Jinja2 sandbox is happier with plain dicts than with
    frozen-slots dataclasses, so the flattening must be total.
    """
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    _write_suggestions(case_dir)

    payload = build_context(case_dir, now=_NOW).as_template_dict()
    # Top-level: dict of dict / scalar
    assert isinstance(payload["case"], dict)
    assert isinstance(payload["suggestions"], dict)
    # Nested suggestion entries are dicts, not dataclasses
    completed = payload["suggestions"]["completed"]
    assert isinstance(completed, list)
    assert isinstance(completed[0], dict)
    assert completed[0]["id"] == "verify-image-hash"


def test_context_tolerates_malformed_suggestions(tmp_path: Path) -> None:
    """A garbage suggestions.json renders as 'no data' rather than raising."""
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    target = caseguide_suggestions_path(case_dir)
    target.parent.mkdir(parents=True)
    target.write_text("{not valid json", encoding="utf-8")

    ctx = build_context(case_dir, now=_NOW)
    assert ctx.suggestions.has_data is False


# --------------------------------------------------------- renderer

pytest.importorskip("docxtpl")
# The renderer tests below only run when docxtpl is installed.

from docx import Document  # noqa: E402  - import after the skip guard

from caseforge.report.cli import main as cli_main  # noqa: E402
from caseforge.report.render import RenderError, render_report  # noqa: E402


def _build_template(target: Path, body: str) -> None:
    """Write a tiny .docx whose only paragraph is ``body``."""
    doc = Document()
    doc.add_paragraph(body)
    doc.save(str(target))


def _read_paragraphs(path: Path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs if p.text]


def test_render_substitutes_top_level_tokens(tmp_path: Path) -> None:
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    template = tmp_path / "template.docx"
    _build_template(
        template,
        "{{ case.name }} ({{ case.reference }}) — {{ examiner.name }}",
    )
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    paragraphs = _read_paragraphs(output)
    assert paragraphs == ["Smoketest CSAM (ST-001) — Alex Smith"]


def test_render_iterates_completed_suggestions(tmp_path: Path) -> None:
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    _write_suggestions(case_dir)
    template = tmp_path / "template.docx"
    _build_template(
        template,
        "{% for s in suggestions.completed %}DONE: {{ s.action }}{% endfor %}",
    )
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    paragraphs = _read_paragraphs(output)
    assert paragraphs == ["DONE: Verify SHA-256."]


def test_render_raises_on_missing_template(tmp_path: Path) -> None:
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    bogus = tmp_path / "does-not-exist.docx"
    with pytest.raises(RenderError, match="Template not found"):
        render_report(
            template_path=bogus,
            context=build_context(case_dir, now=_NOW),
            output_path=tmp_path / "out.docx",
        )


def test_render_raises_on_jinja_syntax_error(tmp_path: Path) -> None:
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    template = tmp_path / "template.docx"
    _build_template(template, "{% for s in suggestions.completed %}oops")  # missing endfor
    with pytest.raises(RenderError):
        render_report(
            template_path=template,
            context=build_context(case_dir, now=_NOW),
            output_path=tmp_path / "out.docx",
        )


# ============================================================ edge cases
#
# Everything below this line was written *after* Phase 1 to actually
# exercise the renderer against scenarios real templates will hit:
# tables, headers/footers, formatting-split tokens, special characters,
# missing optional fields, the CLI entry point, etc.
#


def _build_template_with_table(target: Path, *, header: str, cell: str) -> None:
    """Template with a 2-row, 2-col table — tokens go in cells."""
    doc = Document()
    doc.add_paragraph(header)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Examiner"
    table.cell(1, 1).text = cell
    doc.save(str(target))


def _read_table_cells(path: Path) -> list[list[str]]:
    doc = Document(str(path))
    return [[cell.text for cell in row.cells] for row in doc.tables[0].rows]


def test_render_iterates_table_rows_via_tr_directive(tmp_path: Path) -> None:
    """``{%tr for ... %}`` / ``{%tr endfor %}`` duplicates the rows
    *between* the opener and closer rows. This is the killer feature
    for forensic templates — examiners want one row per evidence item
    or per session in their custody table.

    The two directives must live in **separate** rows (not the same
    row): docxtpl's regex consumes the entire ``<w:tr>`` containing
    a ``{%tr%}`` directive, so a single-row layout would lose the
    closer. Document this gotcha in the test so future readers don't
    re-discover it the hard way.
    """
    case_dir = tmp_path / "case"
    case = Case(
        name="C",
        case_reference="",
        created_at=_NOW,
        updated_at=_NOW,
        scope=ExamScope(evidence_items=["E01 image", "Cellebrite extraction", "Memory dump"]),
    )
    case_dir.mkdir(parents=True)
    write_case(case_dir, case)

    template = tmp_path / "template.docx"
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "#"
    table.cell(0, 1).text = "Item"
    # Row 1: opener (will be removed by docxtpl's row-replacement).
    table.cell(1, 0).text = "{%tr for item in scope.evidence_items %}"
    # Row 2: body (duplicated per item).
    table.cell(2, 0).text = "{{ loop.index }}"
    table.cell(2, 1).text = "{{ item }}"
    # Row 3: closer (will be removed).
    table.cell(3, 0).text = "{%tr endfor %}"
    doc.save(str(template))

    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    rendered = Document(str(output))
    rows = [[cell.text for cell in row.cells] for row in rendered.tables[0].rows]
    assert rows[0] == ["#", "Item"]
    assert len(rows) == 4
    assert rows[1] == ["1", "E01 image"]
    assert rows[2] == ["2", "Cellebrite extraction"]
    assert rows[3] == ["3", "Memory dump"]


def test_render_substitutes_tokens_inside_table_cells(tmp_path: Path) -> None:
    """Tokens in tables are the dominant real-world placement.

    Forensic report templates almost always have a "case header" or
    "scope summary" table with one row per field. If table-cell
    substitution doesn't work, the tool's effectively useless.
    """
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    template = tmp_path / "template.docx"
    _build_template_with_table(
        template, header="Header text", cell="{{ examiner.name }}"
    )
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    cells = _read_table_cells(output)
    assert cells[0] == ["Field", "Value"]
    assert cells[1] == ["Examiner", "Alex Smith"]


def test_render_substitutes_tokens_in_headers_and_footers(tmp_path: Path) -> None:
    """Section headers/footers are where the case reference + page numbers
    typically live; tokens there are the second-most-common placement."""
    case_dir = tmp_path / "case"
    _seed_case(case_dir)

    template = tmp_path / "template.docx"
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Case: {{ case.name }} ({{ case.reference }})"
    section.footer.paragraphs[0].text = "Generated {{ generated_at }}"
    doc.add_paragraph("Body.")
    doc.save(str(template))

    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    rendered = Document(str(output))
    header_text = rendered.sections[0].header.paragraphs[0].text
    footer_text = rendered.sections[0].footer.paragraphs[0].text
    assert "Smoketest CSAM" in header_text
    assert "ST-001" in header_text
    assert "2026-04-26" in footer_text


def test_render_handles_empty_completed_loop_without_error(tmp_path: Path) -> None:
    """A for-block over an empty list should produce no output (not raise)."""
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    # No suggestions written → suggestions.completed is empty.
    template = tmp_path / "template.docx"
    _build_template(
        template,
        "Before {% for s in suggestions.completed %}{{ s.action }}{% endfor %} after",
    )
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    assert _read_paragraphs(output) == ["Before  after"]


def test_render_raises_on_undefined_token(tmp_path: Path) -> None:
    """A typo'd token (``{{ case.bogus }}``) must surface as RenderError.

    docxtpl/Jinja2 raises ``UndefinedError`` for unknown attributes when
    the template runs in strict mode. Our wrapper turns that into
    ``RenderError`` with the class name so the operator sees what kind
    of failure it was.
    """
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    template = tmp_path / "template.docx"
    _build_template(template, "{{ case.bogus_field_that_does_not_exist }}")
    with pytest.raises(RenderError):
        render_report(
            template_path=template,
            context=build_context(case_dir, now=_NOW),
            output_path=tmp_path / "out.docx",
        )


def test_render_preserves_non_ascii_characters(tmp_path: Path) -> None:
    """Examiner names, case names, scope text routinely include non-ASCII
    (smart quotes from copy-paste, accented characters, em-dashes, emoji).
    None of it should round-trip differently than it went in."""
    case_dir = tmp_path / "case"
    case = Case(
        name="Operación Café — résumé",
        case_reference="OP-2026-Δ7",
        created_at=_NOW,
        updated_at=_NOW,
        examiner=ExaminerIdentity(name="Søren O'Hara", organisation="Polícia"),
        scope=ExamScope(exam_type="🔍 fraud"),
    )
    case_dir.mkdir(parents=True)
    write_case(case_dir, case)

    template = tmp_path / "template.docx"
    _build_template(
        template,
        "{{ case.name }} :: {{ examiner.name }} :: {{ scope.exam_type }}",
    )
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    para = _read_paragraphs(output)[0]
    assert "Operación Café — résumé" in para
    assert "Søren O'Hara" in para
    assert "🔍 fraud" in para


def test_render_supports_strftime_on_raw_datetime_fields(tmp_path: Path) -> None:
    """Template authors should be able to format dates however they want.

    The context exposes both pre-formatted strings (``generated_at``)
    AND raw datetime objects (``case.created_at``) so a template can
    render dates in the examiner's preferred format via Jinja's filter
    chaining or attribute access.
    """
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    template = tmp_path / "template.docx"
    # ``%d`` (zero-padded) is the cross-platform spelling — ``%-d`` is
    # POSIX-only and would fail on Windows runners. Examiners on Linux
    # can still write ``%-d`` in their own templates; this test just
    # verifies that strftime is callable on the raw datetime field.
    _build_template(
        template,
        "Created on {{ case.created_at.strftime('%B %d, %Y') }}",
    )
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    # April 26, 2026 — the seed case was created at _NOW (2026-04-26).
    assert _read_paragraphs(output) == ["Created on April 26, 2026"]


def test_unset_custody_received_at_renders_as_empty_string(tmp_path: Path) -> None:
    """``{{ custody.received_at_str }}`` must be empty (not the literal
    string "None") when the case has no custody timestamp.

    Templates that hard-code ``{{ custody.received_at }}`` would print
    "None" otherwise, which looks unprofessional in a finished report.
    The ``_str`` companion field is the safe form to drop into prose.
    """
    case_dir = tmp_path / "case"
    case = Case(
        name="No custody yet",
        case_reference="",
        created_at=_NOW,
        updated_at=_NOW,
        custody=CustodyRecord(),  # received_at=None by default
    )
    case_dir.mkdir(parents=True)
    write_case(case_dir, case)

    template = tmp_path / "template.docx"
    _build_template(template, "Received: [{{ custody.received_at_str }}]")
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    assert _read_paragraphs(output) == ["Received: []"]


def test_set_custody_received_at_renders_via_str_field(tmp_path: Path) -> None:
    """The companion ``received_at_str`` formats the datetime when set."""
    case_dir = tmp_path / "case"
    case = Case(
        name="With custody",
        case_reference="",
        created_at=_NOW,
        updated_at=_NOW,
        custody=CustodyRecord(received_at=_NOW),
    )
    case_dir.mkdir(parents=True)
    write_case(case_dir, case)

    template = tmp_path / "template.docx"
    _build_template(template, "Received: {{ custody.received_at_str }}")
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    assert _read_paragraphs(output) == ["Received: 2026-04-26 12:00 UTC"]


def test_render_with_unset_examiner_renders_empty_strings(tmp_path: Path) -> None:
    """Templates that hard-code examiner tokens still need to render
    cleanly when the case has no examiner identity yet (cases created
    before the examiner walked in). Empty string is correct, not None."""
    case_dir = tmp_path / "case"
    case = Case(
        name="Anonymous case",
        case_reference="",
        created_at=_NOW,
        updated_at=_NOW,
        # ExaminerIdentity defaults are all empty strings.
    )
    case_dir.mkdir(parents=True)
    write_case(case_dir, case)

    template = tmp_path / "template.docx"
    _build_template(
        template,
        "By {{ examiner.name }} ({{ examiner.organisation }})",
    )
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    assert _read_paragraphs(output) == ["By  ()"]


def test_render_iterates_long_suggestion_list(tmp_path: Path) -> None:
    """Real cases can have 50+ suggestions; the iteration shouldn't break
    or silently truncate."""
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    target = caseguide_suggestions_path(case_dir)
    target.parent.mkdir(parents=True, exist_ok=True)
    suggestions = [
        {
            "id": f"step-{i:03d}",
            "action": f"Action {i}.",
            "priority": "required",
            "completed": i % 2 == 0,
        }
        for i in range(50)
    ]
    target.write_text(
        json.dumps(
            {
                "schema_version": 2,
                "scope_summary": "stress",
                "playbooks": [],
                "suggestions": suggestions,
            }
        ),
        encoding="utf-8",
    )

    template = tmp_path / "template.docx"
    _build_template(
        template,
        "{% for s in suggestions.all %}{{ s.id }}|{% endfor %}",
    )
    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    rendered = _read_paragraphs(output)[0]
    assert rendered.count("|") == 50
    assert "step-000|" in rendered
    assert "step-049|" in rendered


def test_render_writes_to_nonexistent_output_directory(tmp_path: Path) -> None:
    """``render_report`` should create intermediate dirs in ``output_path``.

    The CLI accepts arbitrary output paths; the renderer can't assume
    the operator pre-created the directory.
    """
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    template = tmp_path / "template.docx"
    _build_template(template, "{{ case.name }}")
    output = tmp_path / "deep" / "nested" / "out.docx"
    assert not output.parent.exists()
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    assert output.exists()


def test_render_handles_token_split_across_formatting_runs(tmp_path: Path) -> None:
    """Word splits text into multiple "runs" when formatting changes
    mid-paragraph. A naive substituter would see ``{{`` in run 1 and
    ``case.name }}`` in run 2 and fail to match. ``docxtpl`` handles
    this — this test pins that behaviour so a future swap of the
    rendering library doesn't regress."""
    case_dir = tmp_path / "case"
    _seed_case(case_dir)

    template = tmp_path / "template.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    # Three runs, the middle one bold, all forming a single token.
    paragraph.add_run("Case is {{ case")
    bold_run = paragraph.add_run(".")
    bold_run.bold = True
    paragraph.add_run("name }}.")
    doc.save(str(template))

    output = tmp_path / "out.docx"
    render_report(
        template_path=template,
        context=build_context(case_dir, now=_NOW),
        output_path=output,
    )
    text = _read_paragraphs(output)[0]
    assert text == "Case is Smoketest CSAM."


# ---------------------------------------------------------------- CLI


def test_cli_renders_end_to_end(tmp_path: Path) -> None:
    """The ``caseforge-report`` entry point assembles context, renders,
    and exits 0 on success. Important because the CLI is the surface
    examiners script against before the GUI ships."""
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    template = tmp_path / "template.docx"
    _build_template(template, "{{ case.name }} ({{ case.reference }})")
    output = tmp_path / "out.docx"

    code = cli_main(
        [
            "--template",
            str(template),
            "--case",
            str(case_dir),
            "--output",
            str(output),
        ]
    )
    assert code == 0
    assert output.exists()
    assert _read_paragraphs(output) == ["Smoketest CSAM (ST-001)"]


def test_cli_returns_2_for_missing_case(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    _build_template(template, "{{ case.name }}")
    bogus_case = tmp_path / "no-such-case"

    code = cli_main(
        [
            "--template",
            str(template),
            "--case",
            str(bogus_case),
            "--output",
            str(tmp_path / "out.docx"),
        ]
    )
    assert code == 2


def test_cli_returns_3_for_render_failure(tmp_path: Path) -> None:
    case_dir = tmp_path / "case"
    _seed_case(case_dir)
    bogus_template = tmp_path / "missing.docx"

    code = cli_main(
        [
            "--template",
            str(bogus_template),
            "--case",
            str(case_dir),
            "--output",
            str(tmp_path / "out.docx"),
        ]
    )
    assert code == 3
